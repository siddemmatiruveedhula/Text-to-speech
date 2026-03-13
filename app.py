import streamlit as st
import fitz
import os
import re
import tempfile
from google.cloud import texttospeech
from pydub import AudioSegment
from docx import Document
from io import BytesIO

AudioSegment.converter = "ffmpeg"

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "gcp_key.json"

st.set_page_config(page_title="PDF Audio Studio", page_icon="🎧", layout="wide")

st.title("🎧 PDF Audio Studio")
st.caption("Convert PDFs or DOCX files into natural speech using Google Cloud Text-to-Speech")

# ---------------------------
# MODE SELECTOR
# ---------------------------

mode = st.radio(
    "Select Processing Mode",
    ["Upload Files", "Folder Batch Processing"]
)

# ---------------------------
# LOAD VOICES
# ---------------------------

@st.cache_data
def load_voices():

    client = texttospeech.TextToSpeechClient()
    response = client.list_voices()

    voices = []

    for voice in response.voices:

        gender = texttospeech.SsmlVoiceGender(voice.ssml_gender).name

        if "Neural" in voice.name:
            vtype = "Neural"
        elif "Wavenet" in voice.name:
            vtype = "WaveNet"
        elif "Chirp" in voice.name:
            vtype = "Chirp"
        else:
            vtype = "Standard"

        for lang in voice.language_codes:
            voices.append({
                "name": voice.name,
                "language": lang,
                "gender": gender,
                "type": vtype
            })

    return voices


voice_data = load_voices()

# ---------------------------
# SIDEBAR SETTINGS
# ---------------------------

st.sidebar.header("⚙ Voice Settings")

languages = sorted(set(v["language"] for v in voice_data))
selected_language = st.sidebar.selectbox("Language", languages)

genders = sorted(set(v["gender"] for v in voice_data if v["language"] == selected_language))
selected_gender = st.sidebar.selectbox("Gender", genders)

types = sorted(set(v["type"] for v in voice_data if v["language"] == selected_language and v["gender"] == selected_gender))
selected_type = st.sidebar.selectbox("Voice Type", types)

voices = [
    v["name"]
    for v in voice_data
    if v["language"] == selected_language
    and v["gender"] == selected_gender
    and v["type"] == selected_type
]

selected_voice = st.sidebar.selectbox("Voice", voices)

speed = st.sidebar.slider("Speech Speed", 0.5, 2.0, 1.0)
pitch = st.sidebar.slider("Pitch", -10.0, 10.0, 0.0)

client = texttospeech.TextToSpeechClient()

# ---------------------------
# TEXT SPLITTING
# ---------------------------

def split_text(text, max_chars=3000):

    text = re.sub(r'\s+', ' ', text)

    chunks = []
    start = 0

    while start < len(text):

        end = start + max_chars

        if end >= len(text):
            chunks.append(text[start:])
            break

        split_pos = max(
            text.rfind('.', start, end),
            text.rfind('!', start, end),
            text.rfind('?', start, end)
        )

        if split_pos == -1:
            split_pos = text.rfind(' ', start, end)

        if split_pos == -1:
            split_pos = end

        chunks.append(text[start:split_pos].strip())
        start = split_pos + 1

    return chunks

# ---------------------------
# TEXT TO SPEECH
# ---------------------------

def synthesize(text, voice_name):

    chunks = split_text(text)

    audio_segments = []

    progress = st.progress(0)

    for i, chunk in enumerate(chunks):

        synthesis_input = texttospeech.SynthesisInput(text=chunk)

        voice = texttospeech.VoiceSelectionParams(
            language_code="-".join(voice_name.split("-")[0:2]),
            name=voice_name
        )

        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=speed,
            pitch=pitch
        )

        response = client.synthesize_speech(
            input=synthesis_input,
            voice=voice,
            audio_config=audio_config
        )

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        temp_file.write(response.audio_content)
        temp_file.close()

        audio_segments.append(AudioSegment.from_mp3(temp_file.name))

        progress.progress((i + 1) / len(chunks))

    final_audio = sum(audio_segments)

    output_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    final_audio.export(output_file.name, format="mp3")

    return open(output_file.name, "rb").read()

# ---------------------------
# VOICE PREVIEW
# ---------------------------

st.sidebar.markdown("### 🔊 Voice Preview")

if st.sidebar.button("Preview Voice"):

    preview_text = "Hello. This is a preview of the selected voice."

    audio = synthesize(preview_text, selected_voice)

    st.sidebar.audio(audio)

# ---------------------------
# TEXT EXTRACTION
# ---------------------------

def extract_pdf(file):

    # Streamlit uploaded file
    if hasattr(file, "read") and hasattr(file, "name") and not isinstance(file, str):
        data = file.read()
        doc = fitz.open(stream=data, filetype="pdf")

    # Folder file
    else:
        doc = fitz.open(file)

    text = ""

    for page in doc:
        text += page.get_text("text") + "\n"

    return text

def extract_docx(file):

    # If Streamlit uploaded file
    if hasattr(file, "getvalue"):
        doc = Document(BytesIO(file.getvalue()))
    else:
        # Normal file from folder
        doc = Document(file)

    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)

# ---------------------------
# MODE 1 — UPLOAD FILES
# ---------------------------

if mode == "Upload Files":

    uploaded_files = st.file_uploader(
        "Upload PDF or DOCX files",
        type=["pdf", "docx"],
        accept_multiple_files=True
    )

    if uploaded_files:

        if st.button("🎙 Generate Audio"):

            for file in uploaded_files:

                st.subheader(file.name)

                if file.name.endswith(".pdf"):
                    text = extract_pdf(file)

                elif file.name.endswith(".docx"):
                    text = extract_docx(file)

                if not text.strip():
                    st.warning("No readable text found.")
                    continue

                audio = synthesize(text, selected_voice)

                st.audio(audio)

                st.download_button(
                    label="⬇ Download MP3",
                    data=audio,
                    file_name=file.name.rsplit(".", 1)[0] + ".mp3",
                    mime="audio/mp3"
                )

# ---------------------------
# MODE 2 — FOLDER PROCESSING
# ---------------------------

if mode == "Folder Batch Processing":

    st.subheader("Batch Convert Folder")

    input_folder = st.text_input(
        "Enter folder path containing PDF or DOCX files",
        placeholder="C:\TextToSpeech"
    )

    if st.button("🚀 Convert Entire Folder"):

        if not os.path.isdir(input_folder):

            st.error("Invalid folder path")

        else:

            files = [
                f for f in os.listdir(input_folder)
                if os.path.isfile(os.path.join(input_folder, f))
                and f.lower().endswith((".pdf", ".docx"))
            ]

            if not files:

                st.warning("No PDF or DOCX files found")

            else:

                output_folder = os.path.join(input_folder, "audio_output")

                os.makedirs(output_folder, exist_ok=True)

                total_files = len(files)

                # overall progress display
                overall_text = st.empty()

                progress = st.progress(0)

                for i, filename in enumerate(files):

                    current = i+1

                    # show completed/total
                    overall_text.markdown(f"### {current} / {total_files}")

                    # update overall bar
                    progress.progress(current / total_files)

                    st.write(f"Processing: {filename}")
                    file_path = os.path.join(input_folder, filename)

                    with open(file_path, "rb") as f:

                        if filename.endswith(".pdf"):
                            text = extract_pdf(f)

                        elif filename.endswith(".docx"):
                            text = extract_docx(f)

                    if not text.strip():
                        continue

                    audio = synthesize(text, selected_voice)

                    mp3_name = filename.rsplit(".", 1)[0] + ".mp3"

                    mp3_path = os.path.join(output_folder, mp3_name)

                    with open(mp3_path, "wb") as out:
                        out.write(audio)

                    progress.progress((i + 1) / len(files))

                st.success(f"Audio files saved to: {output_folder}")